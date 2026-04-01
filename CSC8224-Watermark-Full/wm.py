#!/usr/bin/env python3
"""
wm.py  —  Watermark CLI Tool
────────────────────────────────────────────────────────────────────────
CSC 8224 Cryptography | Georgia State University | Spring 2026

USAGE
─────
  # Watermark a file (txt or docx):
  python wm.py sign  --file essay.txt       --key "mysecret"
  python wm.py sign  --file report.docx     --key "mysecret"

  # Scan a file — check if it's watermarked (you need the original):
  python wm.py scan  --file received.txt    --original original.txt  --key "mysecret"

  # Scan WITHOUT the original (blind scan — checks only stego layer):
  python wm.py scan  --file received.txt    --key "mysecret"  --blind

  # Batch scan an entire folder:
  python wm.py batch --folder ./documents   --original ./originals  --key "mysecret"

  # Strip watermark from a file:
  python wm.py strip --file watermarked.txt

  # Identify which key watermarked a file (brute-force a key list):
  python wm.py crack --file suspicious.txt  --original original.txt  --keylist keys.txt
────────────────────────────────────────────────────────────────────────
"""

import argparse
import os
import sys
import json
from pathlib import Path

# ── Add parent dir to path so we can import watermark package ────────────────
sys.path.insert(0, str(Path(__file__).parent))

from watermark.core.embedder  import embed, build_watermark, get_metadata
from watermark.core.verifier  import verify
from watermark.core.tokenizer import clean_text
from watermark.core.file_reader import read_file

# ── ZW chars for blind scan ──────────────────────────────────────────────────
ZW_SYNC = '\u2060'
ZW_0    = '\u200b'
ZW_1    = '\u200c'

DIVIDER = "─" * 60


def color(text, code):
    """ANSI color codes for terminal output."""
    codes = {'green': '32', 'red': '31', 'yellow': '33', 'cyan': '36', 'bold': '1'}
    return f"\033[{codes.get(code,'0')}m{text}\033[0m"


# ══════════════════════════════════════════════════════════════════════════════
# COMMAND: sign
# ══════════════════════════════════════════════════════════════════════════════

def cmd_sign(args):
    """Embed a watermark into a file."""
    if not os.path.exists(args.file):
        print(color(f"Error: File not found: {args.file}", 'red'))
        sys.exit(1)

    print(f"\n{DIVIDER}")
    print(color("  🔏  WATERMARK SIGN", 'bold'))
    print(DIVIDER)

    text = read_file(args.file)
    print(f"  File:          {args.file}")
    print(f"  Characters:    {len(text)}")
    print(f"  Words:         {len(text.split())}")
    print(f"  Key (masked):  {'*' * len(args.key)}")

    wm_text = embed(text, args.key, embed_bits=args.bits)
    meta    = json.loads(get_metadata(text, args.key, args.bits))

    # ── Determine output path ─────────────────────────────────────────────────
    base, ext = os.path.splitext(args.file)
    out_path  = args.output or f"{base}_watermarked{ext}"

    if ext.lower() == '.docx':
        from docx import Document
        doc      = Document(args.file)
        wm_words = wm_text.split(' ')
        word_idx = 0
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            para_word_count = len(para.text.split(' '))
            new_text = ' '.join(wm_words[word_idx:word_idx + para_word_count])
            word_idx += para_word_count
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
        doc.save(out_path)
    else:
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(wm_text)

    invisible = sum(1 for c in wm_text if c in (ZW_0, ZW_1, ZW_SYNC))
    print(f"\n  {color('✅ Watermark embedded successfully!', 'green')}")
    print(f"  Output file:     {out_path}")
    print(f"  Document HMAC:   {meta['doc_hmac'][:24]}...{meta['doc_hmac'][-8:]}")
    print(f"  Sentences signed:{meta['sentence_count']}")
    print(f"  Bits embedded:   {meta['embed_bits']} (invisible chars: {invisible})")
    print(f"  Visible changes: {color('NONE — text looks identical', 'green')}\n")


# ══════════════════════════════════════════════════════════════════════════════
# COMMAND: scan
# ══════════════════════════════════════════════════════════════════════════════

def cmd_scan(args):
    """Verify whether a file is watermarked."""
    if not os.path.exists(args.file):
        print(color(f"Error: File not found: {args.file}", 'red'))
        sys.exit(1)

    print(f"\n{DIVIDER}")
    print(color("  🔍  WATERMARK SCAN", 'bold'))
    print(DIVIDER)

    received_text = read_file(args.file)
    print(f"  Scanning:      {args.file}")
    print(f"  Key (masked):  {'*' * len(args.key)}")

    # ── Blind scan (no original needed) ──────────────────────────────────────
    if args.blind or not args.original:
        print(f"  Mode:          {color('BLIND (no original provided)', 'yellow')}")
        print()
        _blind_scan(received_text)
        return

    # ── Full scan (with original) ─────────────────────────────────────────────
    if not os.path.exists(args.original):
        print(color(f"Error: Original file not found: {args.original}", 'red'))
        sys.exit(1)

    original_text = read_file(args.original)
    print(f"  Original:      {args.original}")
    print(f"  Mode:          {color('FULL (with original)', 'cyan')}")
    print()

    vr = verify(received_text, original_text, args.key, args.bits)

    # ── Print verdict ─────────────────────────────────────────────────────────
    verdict_color = 'green' if 'HIGH' in vr.verdict else 'yellow' if 'DEGRADED' in vr.verdict or 'UNCERTAIN' in vr.verdict else 'red'
    print(f"  {color('VERDICT:', 'bold')} {color(vr.verdict, verdict_color)}")
    print(f"  Confidence:    {vr.confidence_score:.4f} / 1.0000")
    print()
    print(f"  Detailed Results:")
    for line in vr.detail:
        print(f"    {line}")

    print()
    _risk_assessment(vr)


def _blind_scan(text: str):
    """Scan for watermark presence without the original or key."""
    has_sync = ZW_SYNC in text
    zw_count = sum(1 for c in text if c in (ZW_0, ZW_1))
    total_zw = sum(1 for c in text if ord(c) in range(0x200B, 0x2070))

    print(f"  Sync marker found:      {color('YES ✓', 'green') if has_sync else color('NO ✗', 'red')}")
    print(f"  Bit characters (0/1):   {zw_count}")
    print(f"  Total zero-width chars: {total_zw}")

    if has_sync and zw_count >= 8:
        print(f"\n  {color('⚠️  WATERMARK SIGNATURE DETECTED', 'yellow')}")
        print(f"  This text likely contains an HMAC-based steganographic watermark.")
        print(f"  To verify authenticity, provide --original and --key.")
    elif total_zw > 0:
        print(f"\n  {color('❓ SUSPICIOUS — zero-width characters found but no sync marker', 'yellow')}")
        print(f"  May have been partially stripped or use a different encoding.")
    else:
        print(f"\n  {color('✅ No watermark payload detected', 'green')}")
        print(f"  Text appears clean. Note: watermark may have been fully stripped.")


def _risk_assessment(vr):
    """Print a human-readable risk assessment based on verification result."""
    print(f"  Risk Assessment:")
    if vr.confidence_score >= 0.85:
        print(color("    ✅ Document is authentic and unmodified.", 'green'))
    elif vr.confidence_score >= 0.50:
        print(color("    ⚠️  Document was watermarked but has been ATTACKED or modified.", 'yellow'))
        if vr.bit_agreement < 0.1:
            print(color("    ⚠️  Steganographic payload was stripped (zero-width chars removed).", 'yellow'))
        if vr.sentence_integrity < 0.8:
            print(color(f"    ⚠️  {vr.sentences_total - vr.sentences_intact} sentence(s) were altered or removed.", 'yellow'))
    elif vr.confidence_score >= 0.20:
        print(color("    ❗ Document heavily modified — possible paraphrase or substitution attack.", 'red'))
        print(color(f"    ❗ Only {vr.jaccard_similarity*100:.0f}% of original vocabulary remains.", 'red'))
    else:
        print(color("    ❌ No evidence of watermark. Either not watermarked or fully rewritten.", 'red'))


# ══════════════════════════════════════════════════════════════════════════════
# COMMAND: batch
# ══════════════════════════════════════════════════════════════════════════════

def cmd_batch(args):
    """Batch scan an entire folder of files."""
    if not os.path.isdir(args.folder):
        print(color(f"Error: Folder not found: {args.folder}", 'red'))
        sys.exit(1)

    print(f"\n{DIVIDER}")
    print(color("  📁  BATCH WATERMARK SCAN", 'bold'))
    print(DIVIDER)
    print(f"  Scanning folder: {args.folder}")
    print()

    extensions = ('.txt', '.docx', '.md')
    files = [
        f for f in Path(args.folder).iterdir()
        if f.suffix.lower() in extensions
    ]

    if not files:
        print(color("  No supported files found (.txt, .docx, .md)", 'yellow'))
        return

    results = []
    for fpath in sorted(files):
        try:
            received_text = read_file(str(fpath))

            # Try to find matching original
            original_text = None
            if args.original and os.path.isdir(args.original):
                orig_path = Path(args.original) / fpath.name
                if orig_path.exists():
                    original_text = read_file(str(orig_path))

            if original_text:
                vr     = verify(received_text, original_text, args.key, args.bits)
                conf   = f"{vr.confidence_score:.3f}"
                status = vr.verdict
            else:
                # Blind scan
                has_sync = ZW_SYNC in received_text
                zw_count = sum(1 for c in received_text if c in (ZW_0, ZW_1))
                conf     = "N/A (blind)"
                status   = "⚠️  PAYLOAD DETECTED" if (has_sync and zw_count >= 8) else "❓ UNKNOWN"

            results.append((fpath.name, conf, status))
            status_color = 'green' if 'HIGH' in status else 'yellow' if any(x in status for x in ['DEGRADED', 'UNCERTAIN', 'DETECTED']) else 'red'
            print(f"  {fpath.name:<35} conf={conf:<12} {color(status, status_color)}")

        except Exception as e:
            print(f"  {fpath.name:<35} {color(f'ERROR: {e}', 'red')}")

    print(f"\n  Scanned {len(files)} file(s).")

    # Save report
    if args.report:
        report_path = args.report
        with open(report_path, 'w') as f:
            f.write("file,confidence,verdict\n")
            for name, conf, status in results:
                f.write(f"{name},{conf},{status}\n")
        print(f"  Report saved: {report_path}")


# ══════════════════════════════════════════════════════════════════════════════
# COMMAND: strip
# ══════════════════════════════════════════════════════════════════════════════

def cmd_strip(args):
    """Strip all zero-width characters from a file."""
    if not os.path.exists(args.file):
        print(color(f"Error: File not found: {args.file}", 'red'))
        sys.exit(1)

    text    = read_file(args.file)
    cleaned = clean_text(text)
    removed = len(text) - len(cleaned)

    base, ext = os.path.splitext(args.file)
    out_path  = args.output or f"{base}_stripped{ext}"

    if ext.lower() == '.docx':
        from docx import Document
        doc      = Document(args.file)
        wm_words = cleaned.split(' ')
        word_idx = 0
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
            para_word_count = len(para.text.split(' '))
            new_text = ' '.join(wm_words[word_idx:word_idx + para_word_count])
            word_idx += para_word_count
            for run in para.runs:
                run.text = ''
            if para.runs:
                para.runs[0].text = new_text
            else:
                para.add_run(new_text)
        doc.save(out_path)
    else:
        with open(out_path, 'w', encoding='utf-8') as f:
            f.write(cleaned)

    print(f"\n  {color('✅ Strip complete', 'green')}")
    print(f"  Removed {removed} zero-width character(s)")
    print(f"  Saved to: {out_path}\n")


# ══════════════════════════════════════════════════════════════════════════════
# COMMAND: crack
# ══════════════════════════════════════════════════════════════════════════════

def cmd_crack(args):
    """Brute-force which key was used to watermark a file."""
    if not os.path.exists(args.file) or not os.path.exists(args.original):
        print(color("Error: --file and --original are required for crack.", 'red'))
        sys.exit(1)
    if not os.path.exists(args.keylist):
        print(color(f"Error: Key list not found: {args.keylist}", 'red'))
        sys.exit(1)

    received = read_file(args.file)
    original = read_file(args.original)

    with open(args.keylist) as f:
        keys = [line.strip() for line in f if line.strip()]

    print(f"\n{DIVIDER}")
    print(color("  🔓  KEY CRACK ATTEMPT", 'bold'))
    print(DIVIDER)
    print(f"  Testing {len(keys)} keys...")
    print()

    best_key   = None
    best_score = 0.0

    for key in keys:
        vr = verify(received, original, key, args.bits)
        if vr.confidence_score > best_score:
            best_score = vr.confidence_score
            best_key   = key
        if vr.is_exact_match:
            print(color(f"  ✅ KEY FOUND: '{key}' (confidence: {vr.confidence_score:.4f})", 'green'))
            return

    print(f"  No exact match found.")
    if best_score > 0.5:
        print(color(f"  Best candidate: '{best_key}' (confidence: {best_score:.4f})", 'yellow'))
    else:
        print(color("  No key produced a convincing match.", 'red'))


# ══════════════════════════════════════════════════════════════════════════════
# ARG PARSER
# ══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        prog='wm',
        description='Cryptographic Text Watermarking CLI — CSC 8224',
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument('--bits', type=int, default=64, help='Bits to embed (default: 64)')
    sub = parser.add_subparsers(dest='command')

    # sign
    p_sign = sub.add_parser('sign', help='Watermark a file')
    p_sign.add_argument('--file',   required=True,  help='Input file (.txt or .docx)')
    p_sign.add_argument('--key',    required=True,  help='Secret key')
    p_sign.add_argument('--output', default=None,   help='Output path (default: <name>_watermarked.<ext>)')

    # scan
    p_scan = sub.add_parser('scan', help='Scan a file for watermark')
    p_scan.add_argument('--file',     required=True, help='File to scan')
    p_scan.add_argument('--key',      required=True, help='Secret key')
    p_scan.add_argument('--original', default=None,  help='Original file for full verification')
    p_scan.add_argument('--blind',    action='store_true', help='Blind scan (no original needed)')

    # batch
    p_batch = sub.add_parser('batch', help='Batch scan a folder')
    p_batch.add_argument('--folder',   required=True, help='Folder to scan')
    p_batch.add_argument('--key',      required=True, help='Secret key')
    p_batch.add_argument('--original', default=None,  help='Folder of original files for comparison')
    p_batch.add_argument('--report',   default=None,  help='Save CSV report to this path')

    # strip
    p_strip = sub.add_parser('strip', help='Strip watermark from file')
    p_strip.add_argument('--file',   required=True, help='File to strip')
    p_strip.add_argument('--output', default=None,  help='Output path')

    # crack
    p_crack = sub.add_parser('crack', help='Brute-force the watermark key')
    p_crack.add_argument('--file',     required=True, help='Suspicious file')
    p_crack.add_argument('--original', required=True, help='Original file')
    p_crack.add_argument('--keylist',  required=True, help='Text file with one key per line')

    args = parser.parse_args()

    if not args.command:
        parser.print_help()
        return

    dispatch = {
        'sign':  cmd_sign,
        'scan':  cmd_scan,
        'batch': cmd_batch,
        'strip': cmd_strip,
        'crack': cmd_crack,
    }
    dispatch[args.command](args)


if __name__ == '__main__':
    main()
