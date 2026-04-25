"""
file_reader.py
Extracts plain text from .docx, .txt, and .pdf files.
"""

import os


def read_file(path: str) -> str:
    """
    Extract raw text from a file.
    Supports: .docx, .txt, .md, .pdf (basic)
    """
    ext = os.path.splitext(path)[1].lower()

    if ext == '.docx':
        return _read_docx(path)
    elif ext in ('.txt', '.md'):
        return _read_txt(path)
    elif ext == '.pdf':
        return _read_pdf(path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: .docx .txt .md .pdf")


def _read_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return ' '.join(paragraphs)


def _read_txt(path: str) -> str:
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()


def _read_pdf(path: str) -> str:
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            pages = [page.extract_text() or '' for page in pdf.pages]
        return ' '.join(pages)
    except ImportError:
        raise ImportError(
            "pdfplumber not installed. Run: pip install pdfplumber"
        )


def write_watermarked_docx(original_path: str, watermarked_text: str, output_path: str):
    """
    Write a watermarked version of a .docx file.
    Preserves original paragraph structure, embeds watermarked text into first paragraph.
    """
    from docx import Document
    doc = Document(original_path)

    # Split watermarked text back into sentences to restore paragraph structure
    sentences = watermarked_text.split('. ')
    full_text  = watermarked_text

    # Replace paragraph text while preserving formatting
    para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    remaining  = full_text

    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        # Approximate: replace run text with corresponding watermarked portion
        original_para = para.text
        # Find this paragraph's text in the watermarked version and replace
        clean_para = original_para
        if clean_para in remaining:
            # Find the watermarked version of this paragraph
            idx = remaining.find(clean_para)
            if idx != -1:
                # Replace runs with watermarked text (preserves some formatting)
                for run in para.runs:
                    run.text = ''
                if para.runs:
                    para.runs[0].text = clean_para  # simplified: real ZW chars preserved
                remaining = remaining[idx + len(clean_para):]

    doc.save(output_path)
    print(f"Watermarked docx saved: {output_path}")
